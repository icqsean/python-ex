import docx

doc = docx.Document("Python開發環境.docx")
print("段落數: ", len(doc.paragraphs))
for p in range(0, 3, 2):
    para = doc.paragraphs[p]
    print("連續文字數: ", len(para.runs))
    for run in para.runs:
        print(run.text)

