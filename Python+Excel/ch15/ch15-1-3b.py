import docx

doc = docx.Document("自動化建立Word文件2.docx")
para = doc.paragraphs[1]
para.insert_paragraph_before("Python是一種直譯語言。")
para1 = doc.add_paragraph("Python可以使用openpyxl套件來自動化" + 
                          "處理Excel的編輯、讀取、建立、儲存" +
                          "、合併儲存格等相關編輯操作。")
para1.add_run("Spyder")
para1.add_run(", ")
para1.add_run("Python IDLE")
doc.save("自動化建立Word文件3.docx")
 
