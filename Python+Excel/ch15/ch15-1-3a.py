import docx

doc = docx.Document("自動化建立Word文件.docx")
doc.add_heading("Python程式設計", level=1)
doc.add_heading("自動化Excel應用", level=2)
doc.save("自動化建立Word文件2.docx")

